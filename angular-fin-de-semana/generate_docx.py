from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Título Principal
title = doc.add_heading('MEMORIA TÉCNICA DESCRIPTIVA PARA REGISTRO DNDA', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Información General
doc.add_heading('1. Información General de la Obra', level=1)
doc.add_paragraph('Título de la Obra: AngularFinDeSemana - Plataforma de Gestión de Inventario Virtual')
doc.add_paragraph('Autor / Titular de los Derechos Patrimoniales: [ESCRIBE TU NOMBRE COMPLETO AQUÍ]')
doc.add_paragraph('Naturaleza de la Obra: Programa de Computador (Software) - Aplicación Web')
doc.add_paragraph('Año de Creación: 2026')

# Descripción General
doc.add_heading('2. Descripción General y Finalidad', level=1)
doc.add_paragraph(
    'AngularFinDeSemana es un sistema de información tipo Single Page Application (SPA) desarrollado bajo '
    'los paradigmas de programación reactiva y arquitectura basada en componentes independientes (Standalone) '
    'de Angular. Su objetivo principal es la gestión, visualización y control de inventarios de '
    'múltiples entidades (Zapatillas, Videojuegos y Cursos) en tiempo real.'
)

# Arquitectura Técnica
doc.add_heading('3. Arquitectura Técnica y Lenguajes', level=1)
doc.add_paragraph('El sistema está construido utilizando el siguiente ecosistema tecnológico:')
doc.add_paragraph('• Lenguaje de Programación: TypeScript (Superconjunto de JavaScript) para la lógica de controladores y servicios.')
doc.add_paragraph('• Markup y Estilos: HTML5 Semántico y CSS3 modular (aislado por componentes).')
doc.add_paragraph('• Framework Frontend: Angular 17+ (Modo Standalone, sin NgModules).')
doc.add_paragraph('• Backend y Persistencia: Firebase Realtime Database integrado a través del SDK oficial @angular/fire.')

# Patrones y Buenas Prácticas
doc.add_heading('4. Patrones de Diseño y Buenas Prácticas de Ingeniería (Clean Code)', level=1)
doc.add_paragraph(
    'El código fuente fue escrito cumpliendo con estándares internacionales de "Clean Code" y los principios SOLID, '
    'específicamente el Principio de Responsabilidad Única (SRP). La lógica de conexión a la base de datos se '
    'encuentra totalmente extraída y encapsulada en Servicios Inyectables (@Injectable).'
)
doc.add_paragraph(
    'Además, se implementaron técnicas avanzadas de optimización de memoria mediante el uso de flujos '
    'de datos asíncronos (Observables) gestionados directamente en las vistas a través del "Async Pipe", '
    'lo que previene fugas de memoria y mejora sustancialmente el rendimiento de la aplicación web.'
)

# Funcionalidades
doc.add_heading('5. Módulos y Funcionalidades Principales', level=1)
doc.add_paragraph('1. Módulo de Autenticación (Login): Control de acceso inicial a la plataforma.')
doc.add_paragraph('2. Panel de Control (Dashboard): Vista central que aloja e instancia los micro-componentes de la SPA.')
doc.add_paragraph('3. Gestión en Tiempo Real (CRUD): Capacidad de inserción (Create), lectura (Read) y eliminación (Delete) de registros (Zapatillas, Videojuegos y Cursos) con sincronización instantánea en la nube, sin requerir recarga de página (Single Page Application).')
doc.add_paragraph('4. Procesamiento de Datos Complejos: Uso de colecciones ES6 (Set) para procesar datos locales y mostrar valores únicos en interfaces de usuario sin iteraciones redundantes.')

# Pruebas Unitarias
doc.add_heading('6. Control de Calidad y Pruebas Unitarias', level=1)
doc.add_paragraph(
    'El código fuente incluye baterías de Pruebas Unitarias (Unit Tests) diseñadas para el motor Jasmine, '
    'empleando dobles de riesgo ("Mocks" y "Spies") para simular la conexión a la base de datos y garantizar '
    'la integridad funcional del software y su escalabilidad a largo plazo.'
)

doc.add_paragraph('\n\n\n__________________________________________________')
doc.add_paragraph('Firma del Autor o Representante Legal')

doc.save('Registro_DNDA_AngularFinDeSemana.docx')
